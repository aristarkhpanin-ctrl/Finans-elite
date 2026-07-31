"""Экспертное заключение и DOCX-документ Финанс-Аудит (фаза E)."""
from __future__ import annotations

from decimal import Decimal
from io import BytesIO

from docx import Document

from app.audit_docgen import build_audit_docx
from audit_core import analyze
from audit_core.models import AuditPeriod, AuditSubjectModel
from audit_core.opinion import build_opinion, opinion_is_positive
from audit_core.samples import build_trading_subject

D = Decimal


def _distressed() -> AuditSubjectModel:
    """Убыточное предприятие с падающей выручкой и вымытым капиталом."""
    return AuditSubjectModel(
        name="ООО «Проблемный»",
        periods=[AuditPeriod(label=y, kind="year") for y in ("2023", "2024")],
        balance={
            "A_FIXED": [D(700), D(700)], "A_INVENTORY": [D(200), D(150)],
            "A_RECEIVABLE": [D(150), D(100)], "A_CASH": [D(80), D(50)],
            "P_EQUITY": [D(430), D(50)], "P_LONG": [D(300), D(200)],
            "P_SHORT": [D(400), D(750)], "M_RETAINED": [D(-50), D(-400)],
        },
        income={
            "I_REVENUE": [D(900), D(600)], "I_COGS": [D(650), D(500)],
            "I_OPEX": [D(230), D(200)], "I_INTEREST": [D(70), D(90)],
            "I_OTHER": [D(0), D(0)], "I_TAX": [D(0), D(0)],
        },
    )


def test_opinion_structure_on_healthy():
    """Заключение содержит интро, показатели, динамику и вердикт."""
    r = analyze(build_trading_subject())
    text = build_opinion(r)
    assert "Анализ фактической отчётности" in text
    assert "По состоянию на последний период (2024)" in text
    assert "коэффициент текущей ликвидности" in text
    assert "За рассматриваемый период" in text and "выручка выросла" in text
    assert text.rstrip().splitlines()[-1].startswith("Итог:")
    # аббревиатуры не ломаются регистром
    assert "(ROE)" in text


def test_opinion_on_distressed_is_negative():
    """Неустойчивое: интро о рисках, перечень нарушений, модели и жёсткий вердикт."""
    r = analyze(_distressed())
    text = build_opinion(r)
    assert "признаки финансовой неустойчивости" in text
    assert "Вне нормативных значений:" in text
    assert "зоне высокого риска" in text
    assert "получен убыток" in text
    assert "неустойчивое" in text.rstrip().splitlines()[-1]
    assert opinion_is_positive(r) is False


def test_opinion_flags_unbalanced_input():
    """Несходящийся баланс явно оговаривается в заключении."""
    m = build_trading_subject()
    m.balance["P_EQUITY"] = [D(5000), D(6000), D(9999)]
    text = build_opinion(analyze(m))
    assert "не сходится" in text and "требуют проверки" in text


def test_opinion_without_periods():
    assert "Недостаточно данных" in build_opinion(analyze(AuditSubjectModel()))


def test_opinion_omits_undefined_metrics():
    """Неопределённые показатели не подставляются нулём, а опускаются."""
    m = build_trading_subject()
    m.balance["P_SHORT"] = [D(0), D(0), D(0)]          # текущая ликвидность не определена
    m.balance["P_EQUITY"] = [D(8000), D(9300), D(11100)]  # баланс сходится
    text = build_opinion(analyze(m))
    assert "коэффициент текущей ликвидности" not in text
    assert "коэффициент автономии" in text


def test_docx_document_structure():
    """DOCX: титул, заключение, отчёты, коэффициенты, диагностика."""
    r = analyze(build_trading_subject())
    content = build_audit_docx(r, build_opinion(r), subject_name="ООО «Торговый дом»",
                               industry="Оптовая торговля", currency="RUB")
    assert content[:2] == b"PK"                        # валидный zip (docx)
    doc = Document(BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "ООО «Торговый дом»" in text
    assert "Заключение по анализу финансового состояния" in text
    assert "Экспертное заключение" in text
    assert "Оптовая торговля" in text

    headings = [p.text for p in doc.paragraphs]
    for expected in ("Баланс (аналитическая форма)", "Отчёт о финансовых результатах",
                     "Коэффициенты — ликвидность", "Диагностика финансового состояния"):
        assert any(expected in h for h in headings), expected
    # таблицы: 2 отчёта + 4 группы коэффициентов + модели + нормативы
    assert len(doc.tables) >= 8


def test_docx_endpoint(client, auth_headers):
    """Эндпоинт документа: DOCX с вложением, изоляция арендатора."""
    sample = {
        "name": "ООО «Тест»", "currency": "RUB", "industry": "Торговля",
        "periods": [{"label": "2024", "kind": "year"}],
        "balance": {"A_FIXED": ["100"], "A_CASH": ["100"],
                    "P_EQUITY": ["150"], "P_SHORT": ["50"]},
        "income": {"I_REVENUE": ["500"], "I_COGS": ["300"], "I_OPEX": ["100"],
                   "I_INTEREST": ["10"], "I_OTHER": ["0"], "I_TAX": ["18"]},
    }
    sid = client.post("/api/v1/audit/subjects", json={"name": "ООО «Тест»", "model": sample},
                      headers=auth_headers).json()["id"]
    r = client.get(f"/api/v1/audit/subjects/{sid}/report.docx", headers=auth_headers)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "attachment" in r.headers["content-disposition"]
    assert r.content[:2] == b"PK"
    assert client.get("/api/v1/audit/subjects/nope/report.docx",
                      headers=auth_headers).status_code == 404


def test_analyze_response_includes_opinion(client, auth_headers):
    sample = {
        "periods": [{"label": "2024", "kind": "year"}],
        "balance": {"A_CASH": ["200"], "P_EQUITY": ["150"], "P_SHORT": ["50"]},
        "income": {"I_REVENUE": ["500"], "I_COGS": ["300"], "I_OPEX": ["100"],
                   "I_INTEREST": ["10"], "I_OTHER": ["0"], "I_TAX": ["18"]},
    }
    sid = client.post("/api/v1/audit/subjects", json={"name": "С", "model": sample},
                      headers=auth_headers).json()["id"]
    body = client.post(f"/api/v1/audit/subjects/{sid}/analyze", headers=auth_headers).json()
    assert body["opinion"].startswith("Анализ фактической отчётности")
    assert "Итог:" in body["opinion"]
