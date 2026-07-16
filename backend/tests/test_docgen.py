"""Тесты DOCX-бизнес-плана (пакет №5, D1): генератор, годовая свёртка, эндпоинт."""
from datetime import date
from decimal import Decimal
from io import BytesIO

from docx import Document

from app.docgen import DOCX_MIME, aggregate_statement, build_business_plan_docx
from calc_core import run
from calc_core.models import PlanSection
from calc_core.reports.lines import CASHFLOW_LINES, PROFIT_USE_LINES
from calc_core.reports.statements import Statement
from calc_core.review import ReviewContext, run_review
from calc_core.review.opinion import build_opinion
from calc_core.samples import build_sample_project


def _build(model=None) -> Document:
    model = model or build_sample_project()
    result = run(model)
    review = run_review(ReviewContext(model=model, result=result))
    content = build_business_plan_docx(model, result, build_opinion(review, result),
                                       project_name="Демо-проект", today=date(2026, 7, 1))
    assert content[:2] == b"PK"                       # валидный zip (docx)
    return Document(BytesIO(content))


def _texts(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def test_document_structure_on_sample():
    doc = _build()
    text = _texts(doc)
    assert "Демо-проект" in text
    assert "Экспертное заключение" in text
    assert "Показатели эффективности инвестиций" in text
    assert "Финансовые отчёты" in text
    assert "Дата формирования: 01.07.2026." in text
    # демо маргинально → заключение с вердиктом о рисках
    assert "не рекомендуется" in text
    # ≥ 5 таблиц: показатели + 4 отчёта
    assert len(doc.tables) >= 5
    # заголовок первого отчёта и метки строк на месте
    stmt = next(t for t in doc.tables if t.rows[0].cells[0].text == "Строка")
    assert any("Чистая прибыль" in r.cells[0].text for r in stmt.rows)


def test_user_sections_in_document():
    model = build_sample_project()
    model.business_plan = [
        PlanSection(title="Резюме проекта", text="Первый абзац.\nВторой абзац."),
        PlanSection(title="", text=""),               # пустой раздел пропускается
    ]
    text = _texts(_build(model))
    assert "Резюме проекта" in text
    assert "Первый абзац." in text and "Второй абзац." in text


def test_monthly_labels_under_limit():
    doc = _build()                                    # демо: 12 мес. с 2026-01-01
    text = _texts(doc)
    assert "Периодичность: помесячно." in text
    stmt = next(t for t in doc.tables if t.rows[0].cells[0].text == "Строка")
    assert stmt.rows[0].cells[1].text == "01.2026"
    assert len(stmt.rows[0].cells) == 13              # «Строка» + 12 месяцев


def test_yearly_aggregation_over_limit():
    model = build_sample_project()
    model.header.duration_months = 30                 # 30 мес. → 3 «года» (последний неполный)
    doc = _build(model)
    text = _texts(doc)
    assert "по годам проекта" in text
    stmt = next(t for t in doc.tables if t.rows[0].cells[0].text == "Строка")
    assert [c.text for c in stmt.rows[0].cells[1:]] == ["Год 1", "Год 2", "Год 3"]


def test_aggregate_statement_rules():
    """Годовая свёртка (Q4): потоки — суммы, C28/P2 — начало, C29/P7 — конец, P3 = P2+ΣP1."""
    n = 24
    cf = Statement(CASHFLOW_LINES, n)
    cf["C1"] = [Decimal(10)] * n                      # поток → суммы по 120
    cf["C28"] = [Decimal(i) for i in range(n)]        # остаток на начало → первый месяц года
    cf["C29"] = [Decimal(i + 1) for i in range(n)]    # остаток на конец → последний месяц года
    agg = aggregate_statement(cf, "flow")
    assert agg["C1"] == [Decimal(120), Decimal(120)]
    assert agg["C28"] == [Decimal(0), Decimal(12)]
    assert agg["C29"] == [Decimal(12), Decimal(24)]

    pu = Statement(PROFIT_USE_LINES, n)
    pu["P1"] = [Decimal(5)] * n
    pu["P2"] = [Decimal(100 + i) for i in range(n)]
    pu["P7"] = [Decimal(200 + i) for i in range(n)]
    agg = aggregate_statement(pu, "flow")
    assert agg["P2"] == [Decimal(100), Decimal(112)]
    assert agg["P7"] == [Decimal(211), Decimal(223)]
    assert agg["P3"] == [Decimal(160), Decimal(172)]  # P2 первого месяца + ΣP1 года

    bal = Statement([("B1", "Деньги")], n)
    bal["B1"] = [Decimal(i) for i in range(n)]
    assert aggregate_statement(bal, "balance")["B1"] == [Decimal(11), Decimal(23)]


def test_yearly_totals_match_monthly_sums():
    """Свёртка реального результата аддитивна: ΣI28 и ΣC13 за год равны годовой строке."""
    model = build_sample_project()
    model.header.duration_months = 36
    result = run(model)
    inc = aggregate_statement(result.income, "flow")
    assert inc["I28"][0] == sum(result.income["I28"][0:12], Decimal(0))
    cf = aggregate_statement(result.cashflow, "flow")
    assert cf["C29"][-1] == result.cashflow["C29"][-1]  # конец последнего года = конец горизонта


# --- эндпоинт ---

def _project(client, headers, name="Документ"):
    sample = client.get("/api/v1/sample").json()
    return client.post("/api/v1/projects", json={"name": name, "model": sample},
                       headers=headers).json()["id"]


def test_business_plan_endpoint(client, auth_headers):
    pid = _project(client, auth_headers)
    r = client.get(f"/api/v1/projects/{pid}/business-plan.docx", headers=auth_headers)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(DOCX_MIME)
    assert "attachment" in r.headers["content-disposition"]
    assert r.content[:2] == b"PK"
    doc = Document(BytesIO(r.content))
    assert any("Экспертное заключение" in p.text for p in doc.paragraphs)


def test_business_plan_missing_and_auth(client, auth_headers):
    assert client.get("/api/v1/projects/nope/business-plan.docx",
                      headers=auth_headers).status_code == 404
    pid = _project(client, auth_headers)
    assert client.get(f"/api/v1/projects/{pid}/business-plan.docx").status_code in (401, 403)
