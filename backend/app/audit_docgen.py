"""Генерация документа заключения по анализу отчётности (Финанс-Аудит, фаза E).

Структура: титул → экспертное заключение → **ядро due diligence** (вердикт, флаги,
качество прибыли, обязательства, оценка, риски, план-факт) → аналитическая форма
(баланс, ОПУ) → коэффициенты по группам → диагностика → чек-лист процедур.
Переиспользует помощники DOCX первого продукта (``_shrink_table``, ``_fmt_money``,
``DOCX_MIME``).

Порядок разделов — порядок чтения инвесткомитетом: сначала вывод и находки, потом
числа, на которых они стоят. Правило состава — SPEC, Приложение У: **в документ идёт
то же, что на экран, вместе со списками «что не посчитано»**. Документ уходит из
системы и читается без экрана; пробел, не названный в нём, читатель примет за
благополучие.
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from typing import Optional

from docx import Document

from audit_core.diagnostics import Diagnostics
from audit_core.earnings import EarningsQuality
from audit_core.flags import FlagRegistry
from audit_core.obligations import ObligationRegister
from audit_core.pipeline import CaseReview
from audit_core.planfact import PlanFact
from audit_core.result import AuditLine
from audit_core.risk import RiskResult
from audit_core.summary import CaseSummary
from audit_core.valuation import Valuation
from calc_core.review.text import fmt_num, fmt_pct

from .docgen import DOCX_MIME, _fmt_money, _shrink_table

__all__ = ["DOCX_MIME", "build_audit_docx"]

_ZONE_LABEL = {"safe": "устойчивость", "grey": "неопределённость", "distress": "высокий риск"}
_STATUS_LABEL = {"good": "норма", "warn": "внимание", "risk": "вне норматива"}
_GROUP_TITLES = [
    ("liquidity", "Ликвидность"),
    ("gearing", "Финансовая устойчивость"),
    ("profitability", "Рентабельность"),
    ("activity", "Деловая активность"),
]
#: Показатели-доли выводятся процентами, денежные — рублями, прочие — числом.
_PCT = ("Рентабельность", "Коэффициент автономии", "Суммарные обязательства")
_MONEY = ("Чистый оборотный капитал",)


def _fmt_ratio(name: str, value: Optional[Decimal]) -> str:
    if value is None:
        return "—"
    if name.startswith(_MONEY):
        return _fmt_money(value)
    if name.startswith(_PCT):
        return fmt_pct(value)
    return fmt_num(value)


def _add_table(doc: Document, title: str, periods: list[str],
               rows: list[tuple[str, list[str], bool]]) -> None:
    """Таблица «строка × период»; третий элемент строки — выделять ли жирным (подытог)."""
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1 + len(rows), cols=1 + len(periods))
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Статья"
    for j, p in enumerate(periods):
        table.rows[0].cells[j + 1].text = p
    for i, (label, values, bold) in enumerate(rows):
        cells = table.rows[i + 1].cells
        cells[0].text = label
        for j, v in enumerate(values):
            cells[j + 1].text = v
        if bold:
            for cell in cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.bold = True
    _shrink_table(table, 8.5)


def _statement_rows(lines: list[AuditLine]) -> list[tuple[str, list[str], bool]]:
    return [(ln.label, [_fmt_money(v) for v in ln.values], ln.subtotal) for ln in lines]


def _add_diagnostics(doc: Document, d: Diagnostics, periods: list[str]) -> None:
    doc.add_heading("Диагностика финансового состояния", level=1)
    doc.add_paragraph(d.summary)
    doc.add_paragraph(f"Оценка выполнена по последнему периоду ({periods[-1]}).")

    scores = [
        (s.name,
         [f"{fmt_num(v)} ({_ZONE_LABEL.get(z, '—')})" if v is not None else "—"
          for v, z in zip(s.values, s.zones, strict=False)],
         False)
        for s in d.scores
    ]
    if scores:
        _add_table(doc, "Модели диагностики банкротства", periods, scores)
        for s in d.scores:
            if s.note:
                doc.add_paragraph(f"{s.name}: {s.note}")

    assessments = [
        (a.name, [_STATUS_LABEL.get(st, "—") if st else "—" for st in a.status], False)
        for a in d.assessments
    ]
    if assessments:
        _add_table(doc, "Оценка показателей по нормативам", periods, assessments)


#: Подписи основ отчётности (совпадают с оговорками свода).
_STANDARD_LABELS = {"rsbu": "РСБУ", "ifrs": "МСФО", "management": "управленческая отчётность"}

#: Подписи серьёзности флага — те же слова, что на экране реестра.
_SEVERITY = {"risk": "Риск", "warning": "Внимание"}


def _not_computed(doc: Document, items: list[str]) -> None:
    """Список «что не посчитано» — обязательная часть раздела, а не сноска.

    Документ читают без экрана: раздел, промолчавший о своём пробеле, читается как
    раздел без пробелов.
    """
    if not items:
        return
    doc.add_paragraph("Не посчитано:")
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_summary(doc: Document, s: CaseSummary) -> None:
    """Вердикт по делу: то же, что в шапке экрана дела (SPEC, Приложение Н)."""
    doc.add_heading("Вердикт по делу", level=1)
    doc.add_paragraph(s.headline)
    if s.detail:
        doc.add_paragraph(s.detail)
    if s.state != "ready":
        _not_computed(doc, s.not_computed)
        return

    line = [f"Флагов риска: {s.risk_flags}", f"предупреждений: {s.warning_flags}"]
    # Оценённое влияние флагов — не скидка к цене (Приложение Н.3), и число
    # неоценённых идёт рядом с суммой всегда: без него сумма читается как полная
    # цена рисков.
    line.append(f"оценённое влияние флагов: {_fmt_money(s.priced_total)}"
                + (f" (без денежной меры ещё {s.unpriced})" if s.unpriced else ""))
    if s.coverage is not None:
        line.append(f"охват проверки: {s.coverage * 100:.0f}%")
    if s.open_procedures:
        line.append(f"незакрытых процедур: {s.open_procedures}")
    if s.input_errors:
        line.append(f"ошибок ввода: {s.input_errors}")
    doc.add_paragraph("; ".join(line) + ".")

    rows = [(m.label, [m.text or (_fmt_money(m.value) if m.unit == "money"
                                  else fmt_num(m.value)) if m.value is not None or m.text
                       else "не считается"],
             False)
            for m in s.metrics]
    if rows:
        _add_table(doc, "Показатели дела", ["Значение"], rows)
    _not_computed(doc, s.not_computed)


def _add_flags(doc: Document, registry: FlagRegistry, periods: list[str]) -> None:
    """Реестр красных флагов. Итог называет две величины и не смешивает их."""
    doc.add_heading("Реестр красных флагов", level=1)
    if not registry.flags:
        doc.add_paragraph("Правила реестра не сработали ни разу: на введённой "
                          "отчётности признаков из каталога не найдено.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells,
                           ("Находка", "Серьёзность", "Периоды", "Оценка влияния"),
                           strict=True):
        cell.text = title
    for f in registry.flags:
        row = table.add_row().cells
        row[0].text = f"{f.title}. {f.detail}" if f.detail else f.title
        row[1].text = _SEVERITY.get(f.severity, f.severity)
        row[2].text = ", ".join(periods[i] for i in f.periods if i < len(periods)) or "—"
        # «Меры нет» — не ноль рублей: флаг без денежного выражения существует,
        # но в сумму не входит.
        row[3].text = _fmt_money(f.impact) if f.impact is not None else "меры нет"
    _shrink_table(table, 8.5)

    tail = (f"Оценённое влияние флагов: {_fmt_money(registry.priced_total)}. "
            "Это не скидка к цене: сумма оценённых находок и торг — разные величины.")
    if registry.unpriced:
        tail += (f" Ещё {registry.unpriced} флагов денежной меры не имеют и в сумму "
                 "не вошли.")
    doc.add_paragraph(tail)


def _add_earnings(doc: Document, e: EarningsQuality, periods: list[str]) -> None:
    """Качество прибыли: отчётный и нормализованный показатель + корректировки."""
    doc.add_heading("Качество прибыли", level=1)
    rows: list[tuple[str, list[str], bool]] = [
        (f"{e.base_code} по отчёту", [_fmt_money(v) for v in e.reported], False)]
    for a in e.adjustments:
        rows.append((f"{a.label} ({a.kind_label})",
                     [_fmt_money(v) for v in a.amounts], False))
    rows.append((f"{e.base_code} нормализованный",
                 [_fmt_money(v) for v in e.normalized], True))
    _add_table(doc, f"Нормализация показателя {e.base_code}", periods, rows)

    if not e.adjustments:
        doc.add_paragraph("Корректировок не вводилось: нормализованный показатель "
                          "равен отчётному.")
    if e.grade:
        note = f"Оценка качества прибыли: {e.grade}"
        if e.deviation is not None:
            note += f" (отклонение от отчётного {fmt_pct(abs(e.deviation))})"
        doc.add_paragraph(note + (f". {e.grade_note}" if e.grade_note else "."))


def _add_obligations(doc: Document, o: ObligationRegister) -> None:
    """Обязательства и залоги. Балансовый долг и забалансовый **не складываются**."""
    doc.add_heading("Обязательства и залоги", level=1)
    doc.add_paragraph(
        f"Долг по реестру (в балансе): {_fmt_money(o.balance_debt)}; "
        f"по балансу отчётности: {_fmt_money(o.reported_debt)}; "
        f"расхождение: {_fmt_money(o.discrepancy)}.")
    # Забалансовые обязательства названы отдельной величиной: сложение с балансовым
    # долгом дало бы число, которого нет ни в одном отчёте.
    doc.add_paragraph(
        f"Забалансовые обязательства (поручительства, гарантии): {_fmt_money(o.off_balance)}. "
        "Отдельная величина: с балансовым долгом она не складывается.")

    if not o.rows:
        doc.add_paragraph("Реестр обязательств не заполнен — сверять нечего.")
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells,
                           ("Кредитор", "Вид", "Сумма", "Срок", "Ковенант"), strict=True):
        cell.text = title
    for r in o.rows:
        cells = table.add_row().cells
        cells[0].text = r.creditor + (f" · {r.contract}" if r.contract else "")
        cells[1].text = r.kind_label + (" (забаланс)" if r.off_balance else "")
        cells[2].text = _fmt_money(r.amount)
        cells[3].text = r.maturity
        cells[4].text = r.covenant_note or r.covenant or "—"
    _shrink_table(table, 8.5)

    pledge = f"Заложено активов: {_fmt_money(o.pledged_total)}"
    if o.pledged_share is not None:
        pledge += f" ({o.pledged_share * 100:.0f}% активов)"
    if o.free_assets is not None:
        pledge += f"; свободных активов: {_fmt_money(o.free_assets)}"
    doc.add_paragraph(pledge + ".")
    if o.covenants_breached or o.covenants_unknown:
        doc.add_paragraph(f"Ковенанты: нарушено {o.covenants_breached}, "
                          f"проверить нечем {o.covenants_unknown}.")


def _add_valuation(doc: Document, v: Valuation) -> None:
    """Оценка стоимости. Не посчитана — названы препятствия, а не ноль."""
    doc.add_heading("Оценка стоимости", level=1)
    if not v.enabled or v.enterprise_value is None:
        doc.add_paragraph("Оценка не посчитана.")
        for b in v.blockers:
            doc.add_paragraph(b, style="List Bullet")
        _not_computed(doc, v.not_computed)
        return

    doc.add_paragraph(
        f"Метод: DCF по свободному денежному потоку фирмы. Ставка дисконтирования "
        f"{fmt_pct(v.wacc)}, темп роста в постпрогнозе {fmt_pct(v.terminal_growth)}, "
        f"база — {v.base_code} {_fmt_money(v.base_ebit)}.")

    if v.years:
        years = [str(y.year) for y in v.years]
        _add_table(doc, "Прогноз свободного денежного потока", years, [
            ("EBIT", [_fmt_money(y.ebit) for y in v.years], False),
            ("Амортизация", [_fmt_money(y.depreciation) for y in v.years], False),
            ("Капвложения", [_fmt_money(y.capex) for y in v.years], False),
            ("Прирост оборотного капитала",
             [_fmt_money(y.nwc_change) for y in v.years], False),
            ("FCFF", [_fmt_money(y.fcff) for y in v.years], True),
            ("Приведённая стоимость",
             [_fmt_money(y.present_value) for y in v.years], False),
        ])

    bridge = [(b.label + (f" — {b.note}" if b.note else ""),
               [_fmt_money(b.amount)], b.kind == "total") for b in v.bridge]
    if bridge:
        _add_table(doc, "Мост от стоимости бизнеса к цене доли", ["Сумма"], bridge)

    tail = [f"Стоимость бизнеса (EV): {_fmt_money(v.enterprise_value)}"]
    if v.terminal_share is not None:
        tail.append(f"доля постпрогноза {fmt_pct(v.terminal_share)}")
    if v.implied_multiple is not None:
        tail.append(f"подразумеваемый мультипликатор {fmt_num(v.implied_multiple)}×")
    if v.equity_value is not None:
        tail.append(f"стоимость доли {_fmt_money(v.equity_value)}")
    if v.equity_min is not None and v.equity_max is not None:
        tail.append(f"диапазон по чувствительности {_fmt_money(v.equity_min)} — "
                    f"{_fmt_money(v.equity_max)}")
    doc.add_paragraph("; ".join(tail) + ".")

    if v.asking_price is not None and v.discount is not None:
        doc.add_paragraph(
            f"Цена продавца: {_fmt_money(v.asking_price)}; дисконт к ней "
            f"{fmt_pct(v.discount)}.")
    for w in v.warnings:
        doc.add_paragraph(w)
    _not_computed(doc, v.not_computed)


def _add_risk(doc: Document, r: RiskResult) -> None:
    """Риски оценки: торнадо и Монте-Карло — с условием их чтения."""
    doc.add_heading("Риски оценки", level=1)
    if not r.available:
        doc.add_paragraph("Анализ рисков не считался.")
        for b in r.blockers:
            doc.add_paragraph(b, style="List Bullet")
        _not_computed(doc, r.not_computed)
        return

    doc.add_paragraph(
        f"Базовая цена доли: {_fmt_money(r.base_price)}. Каждое допущение смещается "
        f"на {fmt_pct(r.step)} в обе стороны.")
    rows = [(t.label,
             [_fmt_money(t.low_price) if t.low_price is not None else "—",
              _fmt_money(t.high_price) if t.high_price is not None else "—",
              _fmt_money(t.span) if t.span is not None else "—"], False)
            for t in r.tornado]
    if rows:
        _add_table(doc, "Чувствительность цены доли",
                   ["Ниже на шаг", "Выше на шаг", "Размах"], rows)

    mc = r.monte_carlo
    if mc is not None:
        doc.add_paragraph(
            f"Монте-Карло: {mc.iterations} прогонов, оценка получена в {mc.valued}"
            + (f", в {mc.unvalued} — нет" if mc.unvalued else "") + ".")
        if mc.median is not None:
            doc.add_paragraph(
                f"Медиана {_fmt_money(mc.median)}; интервал 10–90%: "
                f"{_fmt_money(mc.p10)} — {_fmt_money(mc.p90)}.")
        if mc.below_asking is not None:
            doc.add_paragraph("Доля прогонов с ценой ниже запрошенной: "
                              f"{fmt_pct(mc.below_asking)}.")
        # Условие чтения блока, а не сноска (Приложение Р).
        doc.add_paragraph("Числа Монте-Карло ровно настолько хороши, насколько верны "
                          "заданные распределения допущений: их задаёт аналитик, а не "
                          "платформа.")
    for w in r.warnings:
        doc.add_paragraph(w)
    _not_computed(doc, r.not_computed)


def _add_plan_fact(doc: Document, pf: PlanFact) -> None:
    """План-факт: прогноз продавца против факта (только если план введён)."""
    doc.add_heading("План-факт после сделки", level=1)
    doc.add_paragraph("Сравниваются периоды: " + ", ".join(pf.periods) + ".")
    rows = [(row.label, [_fmt_money(row.plan), _fmt_money(row.fact),
                         _fmt_money(row.delta),
                         fmt_pct(row.delta_share) if row.delta_share is not None else "—"],
             False) for row in pf.rows]
    if rows:
        _add_table(doc, "Прогноз продавца против факта",
                   ["План", "Факт", "Отклонение", "Доля"], rows)

    if pf.flags:
        realized = [f for f in pf.flags if f.realized]
        doc.add_paragraph(
            f"Отмечено сработавшими флагов: {len(realized)} из {len(pf.flags)}. "
            f"Предсказанное влияние (посчитано платформой): "
            f"{_fmt_money(pf.predicted_total)}; фактические потери "
            f"(введены аналитиком): {_fmt_money(pf.realized_total)}."
            + (f" Ещё {pf.unpriced_realized} сработавших флагов не оценены."
               if pf.unpriced_realized else ""))
    for c in pf.caveats:
        doc.add_paragraph(c)
    _not_computed(doc, pf.not_computed)


#: Подписи статусов чек-листа (SPEC, Приложение М.2).
_PROC_STATUS = {
    "pass": "проверено",
    "finding": "есть находка",
    "no_data": "не проверено: нет данных",
    "done": "выполнено аналитиком",
    "skipped": "снято аналитиком",
    "pending": "не выполнено",
}


def _add_procedures(doc: Document, procedures) -> None:
    """Чек-лист процедур и границы проверки (SPEC, Приложение М.4).

    В документ идёт весь каталог, а не только выполненное: отчёт, показывающий одни
    закрытые процедуры, выдаёт часть проверки за целое.
    """
    doc.add_heading("Чек-лист процедур", level=1)
    coverage = procedures.coverage
    head = f"Выполнено процедур: {procedures.closed} из {procedures.total}"
    if coverage is not None:
        head += f" ({coverage * 100:.0f}%)"
    doc.add_paragraph(head + ".")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells,
                           ("Процедура", "Исполнитель", "Итог"), strict=True):
        cell.text = title
    for item in procedures.items:
        row = table.add_row().cells
        row[0].text = item.title
        row[1].text = "платформа" if item.source == "system" else "аналитик"
        status = _PROC_STATUS.get(item.status, item.status)
        row[2].text = f"{status} — {item.detail}" if item.detail else status


def build_audit_docx(review: CaseReview, *, subject_name: str,
                     today: date | None = None) -> bytes:
    """Собрать документ заключения по разбору дела и вернуть содержимое ``.docx``.

    На вход — весь разбор (`review_case`), а не отдельные слои: документ обязан
    рассказывать то же, что экран. Раньше он собирался из результата и чек-листа, а
    находки, оценка и риски до бумаги не доходили — при том, что читатель документа
    и есть тот, кому они адресованы.
    """
    result = review.result
    model = review.model
    industry, currency = model.industry, model.currency
    reporting_standard = model.reporting_standard
    procedures = review.procedures
    doc = Document()

    doc.add_heading(subject_name or "Субъект анализа", level=0)
    doc.add_paragraph("Заключение по анализу финансового состояния")
    meta = [f"Периодов в анализе: {result.n}"]
    if industry:
        meta.append(f"отрасль: {industry}")
    if currency:
        meta.append(f"валюта: {currency}")
    if reporting_standard:
        # Основа отчётности — часть контекста документа: без неё читатель не знает,
        # по каким правилам сформированы статьи, которые он сравнивает.
        meta.append("основа отчётности: "
                    + _STANDARD_LABELS.get(reporting_standard, reporting_standard))
    doc.add_paragraph("; ".join(meta) + ".")
    doc.add_paragraph(f"Дата формирования: {(today or date.today()).strftime('%d.%m.%Y')}.")
    doc.add_paragraph("Финанс-Аудит · анализ по фактической отчётности.")
    if result.revalued:
        # Документ с переоценёнными числами обязан сказать об этом на первой странице.
        doc.add_paragraph("Внимание: показатели рассчитаны по отчётности с учётом "
                          "переоценки статей — они отличаются от учётных данных. "
                          "Перечень поправок приведён в оговорках.")

    doc.add_heading("Экспертное заключение", level=1)
    for block in review.opinion.split("\n\n"):
        for line in block.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # Ядро due diligence — сразу после заключения: вывод и находки читают первыми,
    # числа под ними идут следом.
    _add_summary(doc, review.summary)
    if result.n:
        _add_flags(doc, review.flags, result.periods)
        _add_earnings(doc, review.earnings, result.periods)
        _add_obligations(doc, review.obligations)
        _add_valuation(doc, review.valuation)
        _add_risk(doc, review.risk)
        # План-факт печатается, только когда план введён: раздел «сравнивать не с
        # чем» на бумаге занимает место и ничего не сообщает.
        if review.plan_fact.available:
            _add_plan_fact(doc, review.plan_fact)

    if result.n:
        _add_table(doc, "Баланс (аналитическая форма)", result.periods,
                   _statement_rows(result.balance))
        _add_table(doc, "Отчёт о финансовых результатах", result.periods,
                   _statement_rows(result.income))

        for key, title in _GROUP_TITLES:
            series = result.ratios.get(key) or {}
            if not series:
                continue
            rows = [(name, [_fmt_ratio(name, v) for v in values], False)
                    for name, values in series.items()]
            _add_table(doc, f"Коэффициенты — {title.lower()}", result.periods, rows)

        if result.user_metrics:
            rows = [(u.name + (f" (ошибка: {u.error})" if u.error else ""),
                     [_fmt_money(v) for v in u.values], False)
                    for u in result.user_metrics]
            _add_table(doc, "Пользовательские показатели", result.periods, rows)

        if result.diagnostics is not None:
            _add_diagnostics(doc, result.diagnostics, result.periods)

    if procedures is not None and procedures.total:
        _add_procedures(doc, procedures)

    if not result.balanced:
        doc.add_paragraph("Внимание: введённая отчётность не сходится (актив ≠ пассив).")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
