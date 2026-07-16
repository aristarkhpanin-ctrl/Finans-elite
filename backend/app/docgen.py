"""Генерация DOCX-бизнес-плана (пакет №5, D1; решения Q3–Q4 в BUSINESS-PLAN-DOC-DECOMPOSITION.md).

Документ собирается из готовых модели/результата/заключения — движок не трогаем.
Структура (Q3): титул → экспертное заключение → показатели эффективности →
пользовательские разделы → 4 финансовых отчёта → маржа по продуктам и смета этапов
(при наличии). Длинные горизонты (Q4): при ``n > 24`` отчётные таблицы агрегируются
по годам проекта — свёртка только на этом слое, ядро не меняется.
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt

from calc_core.models import ProjectModel
from calc_core.reports.result import CalcResult
from calc_core.reports.statements import Statement
from calc_core.review.text import fmt_pct, fmt_rub
from calc_core.version import ENGINE_VERSION

#: MIME-тип документа Word (для Response и проверок в тестах).
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#: Порог помесячного отображения (Q4): длиннее — агрегируем по годам проекта.
MONTHLY_LIMIT = 24

# Строки-остатки внутри «потоковых» отчётов: при годовой свёртке берём не сумму,
# а значение на начало (первый месяц года) либо конец (последний месяц года).
_FIRST_OF_PERIOD = {"C28", "P2"}
_LAST_OF_PERIOD = {"C29", "P7"}


def _year_chunks(n: int) -> list[tuple[int, int]]:
    """Полуинтервалы [a, b) месяцев по годам проекта; последний год может быть неполным."""
    return [(a, min(a + 12, n)) for a in range(0, n, 12)]


def _month_labels(start: date, n: int) -> list[str]:
    labels = []
    year, month = start.year, start.month
    for _ in range(n):
        labels.append(f"{month:02d}.{year}")
        month += 1
        if month > 12:
            month, year = 1, year + 1
    return labels


def _year_labels(n: int) -> list[str]:
    return [f"Год {i + 1}" for i in range(len(_year_chunks(n)))]


def aggregate_statement(stmt: Statement, kind: str) -> dict[str, list[Decimal]]:
    """Свернуть отчёт по годам проекта (Q4): потоки — суммы, остатки — начало/конец года.

    ``kind``: ``balance`` — все строки берутся на конец года; иначе строки-остатки
    (C28/P2 — начало, C29/P7 — конец), остальные — суммы. P3 («прибыль к распределению»)
    выводится заново как P2 + P1, чтобы годовой отчёт сохранял тождества помесячного.
    """
    chunks = _year_chunks(stmt.n)
    out: dict[str, list[Decimal]] = {}
    for code in stmt.order:
        series = stmt[code]
        if kind == "balance" or code in _LAST_OF_PERIOD:
            out[code] = [series[b - 1] for _, b in chunks]
        elif code in _FIRST_OF_PERIOD:
            out[code] = [series[a] for a, _ in chunks]
        else:
            out[code] = [sum(series[a:b], Decimal(0)) for a, b in chunks]
    if "P3" in out:
        out["P3"] = [p2 + p1 for p2, p1 in zip(out["P2"], out["P1"], strict=True)]
    return out


def _fmt_money(v: Decimal) -> str:
    return fmt_rub(v)


def _shrink_table(table, size_pt: float = 7.5) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(size_pt)


def _add_statement_table(doc: Document, title: str, stmt: Statement,
                         data: dict[str, list[Decimal]], labels: list[str]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1 + len(stmt.order), cols=1 + len(labels))
    table.style = "Table Grid"
    head = table.rows[0].cells
    head[0].text = "Строка"
    for j, label in enumerate(labels):
        head[j + 1].text = label
    for i, code in enumerate(stmt.order):
        row = table.rows[i + 1].cells
        row[0].text = f"{stmt.labels[code]} ({code})"
        for j, v in enumerate(data[code]):
            row[j + 1].text = _fmt_money(v)
    _shrink_table(table)


def _add_opinion(doc: Document, opinion: str) -> None:
    doc.add_heading("Экспертное заключение", level=1)
    for block in opinion.split("\n\n"):
        for line in block.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())


def _metric_rows(result: CalcResult) -> list[tuple[str, str]]:
    m = result.metrics
    months = lambda v: f"{v} мес." if v is not None else "не достигается"  # noqa: E731
    pct = lambda v: fmt_pct(v) + " годовых" if v is not None else "не определена"  # noqa: E731
    return [
        ("Чистая приведённая стоимость (NPV)", f"{fmt_rub(m.npv)} ₽"),
        ("Внутренняя норма доходности (IRR)", pct(m.irr_annual)),
        ("Модифицированная IRR (MIRR)", pct(m.mirr_annual)),
        ("Средняя норма рентабельности (ARR)", pct(m.arr_annual)),
        ("Индекс прибыльности (PI)", f"{m.pi:.2f}" if m.pi is not None else "не определён"),
        ("Срок окупаемости (PB)", months(m.pb_months)),
        ("Дисконтированный срок окупаемости (DPB)", months(m.dpb_months)),
        ("Приведённая потребность в капитале",
         f"{fmt_rub(m.pv_investments)} ₽" if m.pv_investments is not None else "—"),
        ("Пиковая потребность в финансировании",
         f"{fmt_rub(m.peak_financing_need)} ₽" if m.peak_financing_need is not None else "—"),
    ]


def _add_metrics(doc: Document, result: CalcResult) -> None:
    doc.add_heading("Показатели эффективности инвестиций", level=1)
    rows = _metric_rows(result)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (name, value) in enumerate(rows):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = value
    _shrink_table(table, 9)


def _add_user_sections(doc: Document, model: ProjectModel) -> None:
    for section in model.business_plan:
        if not (section.title or section.text.strip()):
            continue
        doc.add_heading(section.title or "Раздел", level=1)
        for para in section.text.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())


def _add_product_margins(doc: Document, result: CalcResult) -> None:
    pm = result.product_margins
    if not pm.products:
        return
    doc.add_heading("Маржа по продуктам (рецептуры)", level=1)
    table = doc.add_table(rows=1 + len(pm.products), cols=5)
    table.style = "Table Grid"
    for j, h in enumerate(["Продукт", "Выручка", "Себестоимость (BOM + сдельная ЗП)",
                           "Маржа", "Доля маржи"]):
        table.rows[0].cells[j].text = h
    for i, p in enumerate(pm.products):
        row = table.rows[i + 1].cells
        row[0].text = p.name
        row[1].text = _fmt_money(p.revenue)
        row[2].text = _fmt_money(p.bom_cost + p.piece_wages)
        row[3].text = _fmt_money(p.margin)
        row[4].text = fmt_pct(p.margin_share) if p.margin_share is not None else "—"
    _shrink_table(table, 8.5)
    if pm.unallocated_direct:
        doc.add_paragraph(
            f"Нераспределённые (суммовые) прямые издержки: {_fmt_money(pm.unallocated_direct)} ₽."
        )


def _add_budget(doc: Document, result: CalcResult) -> None:
    budget = result.budget
    if not budget.stages:
        return
    doc.add_heading("Смета по этапам календарного плана", level=1)
    kinds = {"expense": "затраты", "asset": "актив", "production": "производство"}
    table = doc.add_table(rows=1 + len(budget.stages), cols=4)
    table.style = "Table Grid"
    for j, h in enumerate(["Этап", "Тип", "Сроки (мес.)", "Стоимость"]):
        table.rows[0].cells[j].text = h
    for i, st in enumerate(budget.stages):
        row = table.rows[i + 1].cells
        row[0].text = st.name
        row[1].text = kinds.get(st.kind, st.kind)
        row[2].text = f"{st.start_month + 1}–{st.finish_month + 1}"
        row[3].text = _fmt_money(st.cost)
    _shrink_table(table, 8.5)
    doc.add_paragraph(f"Итого по смете: {_fmt_money(budget.total)} ₽.")


def _add_statements(doc: Document, model: ProjectModel, result: CalcResult) -> None:
    # Отчётные таблицы широкие → отдельная альбомная секция (Q3/Q4).
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    n = result.n
    yearly = n > MONTHLY_LIMIT
    doc.add_heading("Финансовые отчёты", level=1)
    doc.add_paragraph(
        f"Периодичность: {'по годам проекта (горизонт длиннее 24 мес.)' if yearly else 'помесячно'}."
    )
    if yearly:
        labels = _year_labels(n)
        get = aggregate_statement
    else:
        labels = _month_labels(model.header.start_date, n)
        get = lambda stmt, kind: {code: stmt[code] for code in stmt.order}  # noqa: E731
    _add_statement_table(doc, "Отчёт о прибылях и убытках", result.income,
                         get(result.income, "flow"), labels)
    _add_statement_table(doc, "Кэш-фло", result.cashflow,
                         get(result.cashflow, "flow"), labels)
    _add_statement_table(doc, "Баланс", result.balance,
                         get(result.balance, "balance"), labels)
    _add_statement_table(doc, "Отчёт об использовании прибыли", result.profit_use,
                         get(result.profit_use, "flow"), labels)


def build_business_plan_docx(model: ProjectModel, result: CalcResult, opinion: str,
                             *, project_name: str, today: date | None = None) -> bytes:
    """Собрать документ бизнес-плана (структура Q3) и вернуть содержимое ``.docx``."""
    doc = Document()

    # Титул
    doc.add_heading(project_name or model.header.name, level=0)
    doc.add_paragraph("Бизнес-план")
    doc.add_paragraph(
        f"Горизонт планирования: {result.n} мес. с {model.header.start_date.strftime('%d.%m.%Y')}."
    )
    doc.add_paragraph(f"Дата формирования: {(today or date.today()).strftime('%d.%m.%Y')}.")
    doc.add_paragraph(f"Finans-Elite · движок расчёта v{ENGINE_VERSION}.")

    _add_opinion(doc, opinion)
    _add_metrics(doc, result)
    _add_user_sections(doc, model)
    _add_product_margins(doc, result)
    _add_budget(doc, result)
    _add_statements(doc, model, result)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
